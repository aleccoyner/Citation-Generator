import re
import docx

def authors_function(authors):
    authors = authors.upper()
    authors = authors.split(',')

    idx = authors[0].rfind(" ")
    authors[0] = authors[0][:idx] + ', ' + '.'.join(authors[0][idx+1:]) + '.'

    for i in range (1, len(authors)):
        authors[i] = authors[i][1:]
        idx = authors[i].find(" ")
        authors[i] = '.'.join(authors[i][:idx]) + '.' + authors[i][idx:]

    if len(authors) > 1:
        authors.insert(len(authors) - 1, "AND")

    if len(authors) == 1:
        return authors[0]

    final_string = ""
    for i in range(len(authors)):
        if i < len(authors) - 2:
            final_string += authors[i] + ', '
        elif i == len(authors) - 1 and i != 0:
            final_string += authors[i] + '.'
        else:
            final_string += authors[i] + ' '

    return final_string

input = docx.Document('input.docx')
output = docx.Document()

citations = []
italics = []
for paragraph in input.paragraphs:
    temp = []
    for run in paragraph.runs:
        if run.italic:
            temp.append(run.text.strip())
    italics.append(temp)

    text = paragraph.text.strip()
    if text:
        citations.append(text)

num = 0

for citation in citations:
        new_paragraph = output.add_paragraph()

        split_index = re.search(r"\d{4}\s+", citation).start()
        authors = citation[:split_index].strip()
        citation = citation[split_index:].strip()

        year_index = citation.find("  ") + 1
        year = citation[:year_index - 1]
        citation = citation[year_index:].strip()

        title_index = citation.find("  ") + 1
        title = citation[:title_index - 1]
        citation = citation[title_index:].strip()

        journal_index = citation.find("  ") + 1
        journal = citation[:journal_index - 1]
        citation = citation[journal_index:].strip()

        volume_index = citation.find("  ") + 1
        volume = citation[:volume_index - 1]
        citation = citation[volume_index:].strip()

        pages_index = citation.find("  ") + 1
        if pages_index != 0:
            pages = citation[:pages_index - 1]
            extra = citation[pages_index:].strip()

        else:
            pages = citation
            extra = None

        authors = authors_function(authors)

        new_citation = f'{authors} {year}. {title}. {journal} {volume}:{pages}.'
        if extra:
            new_citation = f'{new_citation} {extra}'

        start_index = 0
        for word in italics[num]:
            if word == '':
                continue
            index = new_citation.find(word, start_index)
            if index != -1:
                new_paragraph.add_run(new_citation[start_index:index])
                runner = new_paragraph.add_run(word)
                runner.italic = True
                start_index = index + len(word)
                index = new_citation.find(word, start_index)

        new_paragraph.add_run(new_citation[start_index:])
        num += 1

output.save('output.docx')
